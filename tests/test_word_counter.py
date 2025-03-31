import pytest
import os
from src.core.word_counter import WordCounter

class TestWordCounter:
    @pytest.mark.unit
    def test_count_words_basic(self):
        """
        Testet die grundlegende Wortzählung.
        """
        # Erstelle eine Testdatei
        test_docx_path = os.path.join('test_documents', 'test_word_count.docx')
        os.makedirs('test_documents', exist_ok=True)
        
        from docx import Document
        doc = Document()
        doc.add_paragraph("Dies ist ein Testdokument mit mehreren Wörtern.")
        doc.save(test_docx_path)
        
        # Führe Wortzählung durch
        result = WordCounter.count_words(test_docx_path)
        
        # Überprüfe Ergebnisse
        assert 'total_words' in result
        assert 'unique_words' in result
        assert 'word_frequency' in result
        
        assert result['total_words'] > 0
        assert result['unique_words'] > 0
        
        # Räume auf
        os.remove(test_docx_path)
    
    @pytest.mark.unit
    def test_count_words_empty_document(self):
        """
        Testet die Wortzählung für ein leeres Dokument.
        """
        # Erstelle eine leere Testdatei
        test_docx_path = os.path.join('test_documents', 'empty_document.docx')
        os.makedirs('test_documents', exist_ok=True)
        
        from docx import Document
        doc = Document()
        doc.save(test_docx_path)
        
        # Führe Wortzählung durch
        result = WordCounter.count_words(test_docx_path)
        
        # Überprüfe Ergebnisse
        assert result['total_words'] == 0
        assert result['unique_words'] == 0
        assert len(result['word_frequency']) == 0
        
        # Räume auf
        os.remove(test_docx_path)
    
    @pytest.mark.unit
    def test_get_word_frequency(self):
        """
        Testet die Wortfrequenz-Funktion.
        """
        words = ['test', 'wort', 'test', 'analyse', 'wort', 'test']
        
        frequency = WordCounter._get_word_frequency(words)
        
        assert len(frequency) > 0
        assert 'test' in frequency
        assert frequency['test'] == 3
        assert frequency['wort'] == 2